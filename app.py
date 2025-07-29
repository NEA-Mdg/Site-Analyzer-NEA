# Importation des bibliothèques 

import streamlit as st
import pandas as pd
import altair as alt
import altair_saver
import numpy as np
import vl_convert as vlc
import plotly.express as px
import plotly.graph_objects as go
import plotly.io as pio
import unicodedata
import matplotlib.pyplot as plt

import io
import re
import os
from io import BytesIO
import tempfile

from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table
from reportlab.lib.styles import getSampleStyleSheet



# Création des fonctions auxiliaires

def clean_statut(x):
    if pd.isna(x) or str(x).strip()=="":
        return "?" 
    x_sans_accent = unicodedata.normalize('NFKD', str(x)).encode('ASCII', 'ignore').decode('utf-8')
    return x_sans_accent.strip().lower()

def sauvegarder_fig_plotly(fig, nom_fichier):
    
    try:
        # Convertir fig Plotly → PNG en mémoire
        img_bytes = fig.to_image(format='png', width=900, height=600, scale=2)
        image = Image.open(io.BytesIO(img_bytes))
        
        # Recréer une fig Matplotlib avec l'image
        fig_mpl, ax = plt.subplots(figsize=(9, 6))
        ax.imshow(image)
        ax.axis('off')
        
        # Sauvegarder l’image temporaire
        with tempfile.TemporaryDirectory() as tmpdirname:
            chemin_complet = os.path.join(tmpdirname, nom_fichier)
            fig_mpl.savefig(chemin_complet, bbox_inches='tight')
            plt.close(fig_mpl)
            # Copier l'image temporaire pour retour stable
            with open(chemin_complet, "rb") as fsrc:
                temp_path = os.path.join(tempfile.gettempdir(), nom_fichier)
                with open(temp_path, "wb") as fdst:
                    fdst.write(fsrc.read())
            return temp_path
    except Exception as e:
        print(f"[Erreur lors de la sauvegarde de la figure : {e}]")
        return None




def generer_rapport_word(site,date_debut, date_fin,date_jour,
                          df_production,img_production, df_etat, img_etat,
                          img1_evolution,img2_evolution,inclure_prod=True,inclure_etat=True, inclure_evolution=True,
                          inclure_synthese_prod=True,inclure_repartition_prod=True, inclure_etat_dominant=True,inclure_etat_repartition=True,inclure_prod_solaire=True,inclure_prod_source=True,logo_path=None):
    doc = Document()

    def add_text_paragraph(text,bold=False, italic=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.bold = bold
        run.italic = italic
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
        run.font.color.rgb = RGBColor(0, 0, 0)

    def add_heading2(text):
        p = doc.add_heading(text, level=1)
        run = p.runs[0]
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 51, 102)  # bleu foncé

    
    def add_centered_plotly_image(doc, image_path, width_in_inches=5):
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run()
        if not image_path or not os.path.exists(image_path):
            run.add_text("[Erreur: figure absente]")
            return
        try:
            run.add_picture(image_path, width=Inches(width_in_inches))
        except Exception as e:
            run.add_text(f"[Erreur: impossible d’ajouter le graphique – {e}]")
                
    def add_table_from_df(df, afficher_index=True):
        df = df.copy()
        if afficher_index:
        
            if df.index.name is None:
                df.index.name = ""
            df_reset = df.reset_index()
        else:
            df_reset = df

        cols = len(df_reset.columns)
        table = doc.add_table(rows=1, cols=cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        for row in table.rows:
            for cell in row.cells:
                cell.width = Inches(1.75)

       


        # Entête
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df_reset.columns):
            paragraph = hdr_cells[i].paragraphs[0]
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph.style = doc.styles['Normal']
            run = paragraph.add_run(str(col))
            run.font.name = 'Calibri'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
            run.font.size = Pt(10)

        # Données
        for _, row in df_reset.iterrows():
            row_cells = table.add_row().cells
            for i, item in enumerate(row):
                paragraph = row_cells[i].paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                paragraph.style = doc.styles['Normal']
                run = paragraph.add_run(str(item) if pd.notna(item) else "—")
                run.font.name = 'Calibri'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
                run.font.size = Pt(10)

    
    # Logo en haut de page
    if logo_path:
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run_logo = p_logo.add_run()
        run_logo.add_picture(logo_path, width=Inches(1))
    
    # Titre principal
    titre = doc.add_heading(f"Rapport d'analyse {site}", 0)
    titre.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = titre.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.bold = True

    # Infos générales
    add_text_paragraph(f"Période analysée : {date_debut} → {date_fin}")

    # ========================
    # 1 - PRODUCTION ENERGETIQUE
    # ========================
    if inclure_prod :
        add_heading2("Production énergétique")
     

        if inclure_synthese_prod :
            # Tableau de production
            add_text_paragraph("Synthèse de production par source",bold=True)
            add_table_from_df(df_production,afficher_index=True)

        if inclure_repartition_prod :
            doc.add_paragraph()
            # Graphe camembert
            add_text_paragraph("Répartition de la production par source",bold=True)
            add_centered_plotly_image(doc,img_production)
        
        doc.add_page_break()

    # ========================
    # 2 - ETAT DE FONCTIONNEMENT
    # ========================

    if inclure_etat :
        add_heading2("État de fonctionnement")
    

        if inclure_etat_dominant:
            # Tableau état dominant par source
            add_text_paragraph("État dominant par source",bold=True)
            add_table_from_df(df_etat,afficher_index=False)

        if inclure_etat_repartition :
            doc.add_paragraph()
            # Graphe camembert
            add_text_paragraph("Répartition de l’état de l’installation globale",bold=True)
            add_centered_plotly_image(doc,img_etat)
        doc.add_page_break()

    # ========================
    # 3 - EVOLUTION TEMPORELLE
    # ========================

    if inclure_evolution :
        
        add_heading2("Évolution temporelle")
     

        if inclure_prod_solaire:

            # Graphique production réelle vs théorique
            add_text_paragraph("Production solaire réelle vs théorique (Energie)",bold=True)
            add_centered_plotly_image(doc,img1_evolution)

        if inclure_prod_source :
            doc.add_paragraph()
            # Graphique production quotidienne par source
            add_text_paragraph("Production quotidienne par source (Puissance)",bold=True)
            add_text_paragraph(f"Résultat du {date_jour}",italic=True)
            add_centered_plotly_image(doc,img2_evolution)

    temp_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_path.name)
    return temp_path.name
    


# Configuration de l'affichage outil
st.set_page_config(
	page_title="Site Analyzer", 
	page_icon="🔆",
	layout="wide",
	initial_sidebar_state="expanded"
	)


# Initialisation des variables
if "site_name" not in st.session_state:
    st.session_state.site_name = ""
if "fichier_donnees" not in st.session_state:
    st.session_state.fichier_donnees = None
if "df_donnees" not in st.session_state:
    st.session_state.df_donnees = None



# Barre latérale pour la navigation
onglet = st.sidebar.radio(
    "",
    ["💡 Indications", "📁 Chargement de données", "📊 Analyse & Visualisation"]
)



# Onglet 1 
if onglet == "💡 Indications":
    st.title("💡 Indications")

    
    st.markdown("---")

    st.subheader("📄 Format du fichier attendu")


    colonnes_attendues = [
        'date', 'heure',
        'puissance_grid', 'puissance_ge', 'puissance_solaire', 'puissance_conso',
        'energie_grid', 'energie_ge', 'energie_solaire', 'energie_solaire_theorique',
        'energie_conso',
        'statut_grid', 'statut_ge', 'statut_solaire', 'statut_installation'
    ]

    df_exemple = pd.DataFrame([
        ['2025-07-10', '08:00', 5.2, 1.3, 3.4, 9.9, 1.1, 0.5, 2.8, 3.0, 4.5, 'on', 'normal', 'excellent', 'ras'],
        ['2025-08-18', '08:10', 5.3, 1.4, 3.6, 10.1, 1.2, 0.6, 2.9, 3.2, 4.7, 'on', 'sous-regime', 'critique', 'panne nea'],
        ['2025-10-10', '13:20', 5.4, 1.2, 3.7, 10.3, 1.3, 0.4, 3.0, 3.3, 4.8, 'off', 'eteint', 'mauvais', 'ecretage client'],
        ['2025-12-03', '18:00', 5.4, 1.2, 3.7, 10.3, 1.3, 0.4, 3.0, 3.3, 4.8, 'off', 'eteint', 'tolerable', 'ecretage client']
    ], columns=colonnes_attendues)

    st.dataframe(df_exemple)

    st.markdown("""
    #### 📌 Détails des colonnes

    - `date` : date au format `YYYY-MM-DD`
    - `heure` : heure au format `HH:MM` 
    - `puissance_*` : puissances en **kW**
    - `energie_*` : énergies en **kWh**
    - `statut_grid` : `"on"` ou `"off"`
    - `statut_ge` : `"eteint"`, `"normal"` ou `"sous-regime"`
    - `statut_solaire` : `"critique"`, `"mauvais"`, `"tolerable"` ou `"excellent"`
    - `statut_installation` : `"panne nea"`, `"ecretage client"`, `"ras"` ou `""` (vide)

    *Le fichier doit contenir **exactement 15 colonnes** dans cet ordre précis.*
    """)

    st.markdown("---")
    st.markdown(" ⚠️ Points de vigilance : ")
    st.warning("""
    - 1/ Les fichiers doivent être au **format Excel (.xlsx / .xlsm) ou CSV (.csv)**  
    - 2/ Respecter strictement l’**ordre des colonnes** tel que défini pour le tableau de données (les noms ne sont pas importants)
    - 3/ Vérifier que le **format de date et heure** utilisé soit bien conforme à celui requis 
    - 4/ S’assurer que toutes les valeurs soient exprimées dans les **unités** demandées
    - 5/ Vérifier que toutes les **valeurs numériques et textuelles** soient **valides** (aucun caractère non-autorisé ne doit être présent )
    """)    

    st.markdown("A.R.")

# Onglet 2 
elif onglet == "📁 Chargement de données":
    st.title("📁 Chargement de données")

    # 1. Entrée du nom du site
    site_name = st.text_input("Nom du site", value=st.session_state.site_name)
    st.session_state.site_name = site_name

    st.write("")

    # 2. Upload du fichier unique
    fichier = st.file_uploader("📄 Importer le fichier de données (CSV ou Excel)", type=["csv", "xlsx"])

    st.write("")
    st.write("")

    # 3. Bouton de validation
    if st.button("Valider"):

        if fichier is None:
            st.error("❌ Aucun fichier n’a été importé.")
        else:
            try:
                if fichier.name.endswith(".csv"):
                    df = pd.read_csv(fichier)
                else:
                    df = pd.read_excel(fichier)

                if df.shape[1] != 15:
                    st.error(f"❌ Le fichier contient {df.shape[1]} colonnes au lieu de 15.")
                else:
                    # Sauvegarde en session
                    st.session_state.fichier_donnees = fichier
                    st.session_state.df_donnees = df
                    st.success(f"✅ Données du site {site_name} chargées avec succès ({df.shape[0]} lignes, 15 colonnes).")

            except Exception as e:
                st.error(f"❌ Erreur lors de la lecture du fichier : {e}")

# Onglet 3
elif onglet == "📊 Analyse & Visualisation":
    st.title(f"📊 Analyse & Visualisation {st.session_state.site_name}")

    if st.session_state.df_donnees is None:
        st.warning("⚠️ Aucune donnée chargée. Veuillez d’abord importer un fichier dans l’onglet précédent.")
    else:
        df = st.session_state.df_donnees.copy()

        #-----------------------
        # NETTOYAGE DE DONNEES
        #-----------------------

        # 1. Renommage des colonnes
        df.columns = [
            'date', 'heure',
            'puissance_grid', 'puissance_ge', 'puissance_solaire', 'puissance_conso',
            'energie_grid', 'energie_ge', 'energie_solaire', 'energie_solaire_theorique',
            'energie_conso',
            'statut_grid', 'statut_ge', 'statut_solaire', 'statut_installation'
        ]

        # 2. Création d'une colonne datetime (fusion date + heure)
        df['datetime'] = pd.to_datetime(df['date'] + ' ' + df['heure'], format='%Y-%m-%d %H:%M:%S')

        # 3. Conversion des colonnes numériques en float
        cols_numeriques = [
            'puissance_grid', 'puissance_ge', 'puissance_solaire', 'puissance_conso',
            'energie_grid', 'energie_ge', 'energie_solaire', 'energie_solaire_theorique',
            'energie_conso'
        ]

        for col in cols_numeriques:
            df[col] = pd.to_numeric(df[col], errors='coerce')

        # 4. Nettoyage des colonnes de statut
        cols_statut = ['statut_grid', 'statut_ge', 'statut_solaire', 'statut_installation']

        for col in cols_statut:
            df[col] = df[col].apply(clean_statut)

        # 5. Réorganisation des colonnes : mettre datetime en premier
        colonnes_ordre = ['datetime'] + [col for col in df.columns if col != 'datetime']
        df = df[colonnes_ordre]
        df = df.sort_values("datetime").reset_index(drop=True)

        #-----------------------
        # TRAITEMENT DE DONNEES
        #-----------------------

        # 1. Choix de la période d'analyse

        st.write("")
        st.write("")

        st.write("**📅 Sélection de la période d’analyse**")

        st.write("*Pour une analyse sur un jour, sélectionnez la **même date** en début et fin.*")

        min_date = df["datetime"].min().date()
        max_date = df["datetime"].max().date()

        col1, col2 = st.columns(2)
        with col1:
            date_debut = st.date_input("**Date de début**", min_value=min_date, max_value=max_date, value=min_date)
        with col2:
            date_fin = st.date_input("**Date de fin**", min_value=min_date, max_value=max_date, value=max_date)
        
        if date_fin < date_debut:
            st.error("❌ La date de fin doit être supérieure ou égale à la date de début.")
            st.stop()
        
        df_data= df[(df["datetime"].dt.date >= date_debut) & (df["datetime"].dt.date <= date_fin)]

        
        st.write("")
        st.write("")
        st.write("")

        # 2. Production énergétique

        st.header("🔋 Production énergétique")

        st.write("")

        # >>>> Tableau synthèse

        st.markdown("**🔍Synthèse de production par source**")

        # Calcul des données utiles
        pic_grid = df_data["puissance_grid"].max()
        pic_ge = df_data["puissance_ge"].max()
        pic_solaire = df_data["puissance_solaire"].max()
        pic_total = df_data["puissance_conso"].max()

        energie_grid = df_data["energie_grid"].sum()
        energie_ge = df_data["energie_ge"].sum()
        energie_solaire = df_data["energie_solaire"].sum()
        energie_solaire_theo = df_data["energie_solaire_theorique"].sum()
        energie_totale = df_data["energie_conso"].sum()

        pertes_solaire = energie_solaire_theo - energie_solaire

        pas_min = 10
        h_marche_grid = df_data[df_data["puissance_grid"] > 0].shape[0] * pas_min / 60
        h_marche_ge = df_data[df_data["puissance_ge"] > 0].shape[0] * pas_min / 60
        h_marche_solaire = df_data[df_data["puissance_solaire"] > 0].shape[0] * pas_min / 60

        # Tableau croisé de synthèse
        tableau1 = pd.DataFrame({
            "Grid": [
                round(pic_grid, 2),
                round(h_marche_grid, 2),
                round(energie_grid, 2),
                None,
                None
            ],
            "GE": [
                round(pic_ge, 2),
                round(h_marche_ge, 2),
                round(energie_ge, 2),
                None,
                None
            ],
            "Solaire": [
                round(pic_solaire, 2),
                round(h_marche_solaire, 2),
                round(energie_solaire, 2),
                round(energie_solaire_theo, 2),
                round(pertes_solaire, 2)
            ],
            "Installation globale": [
                round(pic_total, 2),
                None,
                round(energie_totale, 2),
                None,
                None
            ]
        }, index=[
            "Pic de puissance (kW)",
            "Heures de marche (h)",
            "Énergie réelle produite (kWh)",
            "Énergie théorique produite (kWh)",
            "Pertes énergétiques (kWh)"
        ])

        st.dataframe(tableau1.style.format(na_rep="—"), use_container_width=True)

        st.write("")
        st.write("")

        # >>>> Répartition énergétique

        st.markdown("**🔍 Répartition de la production totale**")

        # Création des données pour le camembert
        labels = ["Grid", "GE", "Solaire"]
        values = [energie_grid, energie_ge, energie_solaire]
        colors = ["#8B2A03", "#003366", "#FFA500"]  

        # Création du graphique avec Plotly
        fig = go.Figure(data=[go.Pie(
        labels=labels,
        values=values,
        marker=dict(colors=colors),
        )])
        
        # Affichage du graphique
        fig.update_layout(title_text="")
        st.plotly_chart(fig, use_container_width=True)

        st.write("")
        st.write("")
        st.write("")
        
        # 3. Etat de fonctionnement

        st.header("🛠️ Etat de fonctionnement")
        
        st.write("")
        
        # >>>> Etat dominant par source
        st.markdown("**🔍 État dominant par source**")

        # Normalisation des statuts 
        df_data["statut_solaire"] = df_data["statut_solaire"].replace("mauvaise", "mauvais")

        # Remplacer les valeurs manquantes par 'absence de données'
        df_data["statut_grid"] = df_data["statut_grid"].fillna("?")
        df_data["statut_ge"] = df_data["statut_ge"].fillna("?")
        df_data["statut_solaire"] = df_data["statut_solaire"].fillna("?")
        df_data["statut_installation"] = df_data["statut_installation"].fillna("?")

        # Définition d'une fonction qui trouve la valeur la plus fréquente
        def statut_dominant(colonne, ignorer_eteint=False):
            serie = df_data[colonne]
            
            if ignorer_eteint and colonne == "statut_ge":
                serie = serie[serie != "eteint"]
            
            mode_val = serie.mode()
            return mode_val.iloc[0] if not mode_val.empty else "?"

        etat_dominant = {
            "Grid": statut_dominant("statut_grid"),
            "GE": statut_dominant("statut_ge",ignorer_eteint=True),
            "Solaire": statut_dominant("statut_solaire"),
            "Installation globale": statut_dominant("statut_installation")
        }

        # Création du DataFrame pour l'affichage
        df_etat_dominant = pd.DataFrame(etat_dominant, index=["Statut dominant"]).T.reset_index()
        df_etat_dominant.columns = ["Source", "Statut dominant"]

        # Affichage du tableau
        st.dataframe(df_etat_dominant, use_container_width=True, hide_index=True)

        st.write("")
        st.write("")
        
        # >>>> Répartition de l'état de l’installation globale
        st.markdown("**🔍 Répartition de l’état de l’installation globale**")

        # Comptage des occurrences
        repartition_etat = df_data["statut_installation"].value_counts().reset_index()
        repartition_etat.columns = ["Statut", "Nombre"]
 
        # Personnalisation des couleurs personnalisées 
        couleurs_etats = {
            "panne nea": "#D62728",           
            "ecretage client": "#0F58DF",     
            "ras": "#2CA02C",                 
            "?": "#B0B0B0"   
        }

        # Association d' une couleur à chaque statut présent
        repartition_etat["Couleur"] = repartition_etat["Statut"].map(couleurs_etats)

        # Création du camembert 
        fig_etat = go.Figure(
            data=[
                go.Pie(
                    labels=repartition_etat["Statut"],
                    values=repartition_etat["Nombre"],
                    marker=dict(colors=repartition_etat["Couleur"]),
                    textinfo="percent"
                )
            ]
        )

        # Affichage du graphique
        st.plotly_chart(fig_etat, use_container_width=True)


        st.write("")
        st.write("")
        st.write("")
        
        # 3. Evolution temporelle

        st.header("📈 Évolution temporelle")
        
        st.write("")

        # >>>> Production solaire réelle vs théoriquee 
        st.markdown("**🔍 Production solaire réelle vs théorique (Énergie)**")
                
        # Préparation des données : grouper par heure (sur 24h)
        df_energy = df_data.copy()
        df_energy["heure"] = df_energy["datetime"].dt.strftime("%H")
        df_energy_grouped = df_energy.groupby("heure")[["energie_solaire", "energie_solaire_theorique"]].sum().reset_index()

        # Création du graphique
        fig1 = go.Figure()

        # Barres : énergie solaire réelle
        fig1.add_trace(go.Bar(
            x=df_energy_grouped["heure"],
            y=df_energy_grouped["energie_solaire"],
            name="E. solaire réelle (kWh)",
            marker_color="#FFA500",
            hovertemplate="Heure : %{x}<br>Energie solaire réelle : %{y:.2f} kWh<extra></extra>"
        ))

        # Courbe : énergie solaire théorique
        fig1.add_trace(go.Scatter(
            x=df_energy_grouped["heure"],
            y=df_energy_grouped["energie_solaire_theorique"],
            name="E. solaire théorique (kWh)",
            mode="lines+markers",
            line=dict(color="#EC0E0E", width=3),
            hovertemplate="Heure : %{x}<br>Energie solaire théorique : %{y:.2f} kWh<extra></extra>"
        ))

        # Mise en forme
        fig1.update_layout(
            xaxis_title="Heure de la journée",
            yaxis_title="Énergie (kWh)",
            barmode="group",
            template="simple_white",
            
        )


        st.plotly_chart(fig1, use_container_width=True)

        st.write("")
        st.write("")

        # >>>> Production quotidienne par source 

        st.markdown("**🔍 Production quotidienne par source (Puissance)**")

        st.write("")

        
        # Filtrage des jours disponibles dans la période sélectionnée
        jours_disponibles = df_data["datetime"].dt.date.unique()

        # Selection pour choisir un jour
        jour_choisi = st.selectbox("📆 Choisir un jour", options=jours_disponibles)

        # Filtrage des données du jour choisi
        df_jour = df_data[df_data["datetime"].dt.date == jour_choisi].copy()

        # Extraction de l'heure exacte pour affichage précis
        
        df_jour["heure"] = df_jour["datetime"].dt.strftime("%H:%M")
        

        # Création du graphique
        fig2 = go.Figure()

        fig2.add_trace(go.Scatter(
            x=df_jour["heure"],
            y=df_jour["puissance_grid"],
            mode="lines",
            name="Grid",
            line=dict(color="#8B2A03")
        ))

        fig2.add_trace(go.Scatter(
            x=df_jour["heure"],
            y=df_jour["puissance_ge"],
            mode="lines",
            name="GE",
            line=dict(color="#003366")
        ))

        fig2.add_trace(go.Scatter(
            x=df_jour["heure"],
            y=df_jour["puissance_solaire"],
            mode="lines",
            name="Solaire",
            line=dict(color="#FFA500")
        ))

        fig2.add_trace(go.Scatter(
            x=df_jour["heure"],
            y=df_jour["puissance_conso"],
            mode="lines",
            name="Installation globale",
            line=dict(color="#6B6767", dash="dash")
        ))

        # Mise en forme
        fig2.update_layout(
            xaxis_title="Heure",
            yaxis_title="Puissance (kW)",
            template="simple_white",
          
        )



        st.plotly_chart(fig2, use_container_width=True)


        st.write("")
        st.write("")
        st.write("")
        st.write("")
        st.write("")




        #-----------------------
        # Génération de rapport
        #-----------------------

        st.markdown("### 📑 Sections à inclure dans le rapport")

        # Sections principales
        inclure_prod = st.checkbox("**1. Production énergétique**", value=True)
        inclure_synthese_prod = st.checkbox("↳ Synthèse de production par source", value=True)
        inclure_repartition_prod = st.checkbox("↳ Répartition de la production totale", value=True)

        inclure_etat = st.checkbox("**2. État de fonctionnement**", value=True)
        inclure_etat_dominant = st.checkbox("↳ État dominant par source", value=True)
        inclure_etat_repartition = st.checkbox("↳ Répartition de l’état de l'installation globale", value=True)

        inclure_evolution = st.checkbox("**3. Évolution temporelle**", value=True)
        inclure_prod_solaire = st.checkbox("↳ Production solaire réelle vs théorique (Energie)", value=True)
        inclure_prod_source = st.checkbox("↳ Production quotidienne par source (Puissance)", value=True)

        # Bouton de génération
        if st.button("Générer le rapport"):

            # === Paramètres à récupérer dynamiquement ===
            site = st.session_state.site_name  # à adapter selon ton app
            date_debut = date_debut.strftime("%Y-%m-%d")
            date_fin = date_fin.strftime("%Y-%m-%d")
            date_jour= jour_choisi.strftime("%Y-%m-%d")

            img_prod_path = sauvegarder_fig_plotly(fig, "repartition_production.png")
            img_etat_path = sauvegarder_fig_plotly(fig_etat, "repartition_etat.png")
            img_ev1_path = sauvegarder_fig_plotly(fig1, "prod_reelle_vs_theorique.png")
            img_ev2_path = sauvegarder_fig_plotly(fig2, "prod_quotidienne_sources.png")


            # === Appel de la fonction de génération ===
            rapport_path = generer_rapport_word (
            site=st.session_state.site_name,
            date_debut=date_debut,
            date_fin=date_fin,
            date_jour=jour_choisi,
            df_production=tableau1,
            img_production=img_prod_path,
            df_etat=df_etat_dominant,
            img_etat=img_etat_path,
            img1_evolution=img_ev1_path,
            img2_evolution=img_ev2_path,
            inclure_prod=inclure_prod,
            inclure_etat=inclure_etat,
            inclure_evolution=inclure_evolution,
            inclure_synthese_prod=inclure_synthese_prod,
            inclure_repartition_prod=inclure_repartition_prod,
            inclure_etat_dominant=inclure_etat_dominant,
            inclure_prod_solaire=inclure_prod_solaire,
            inclure_prod_source=inclure_prod_source,
            logo_path="logo_NEA.png"
            )
            with open(rapport_path, "rb") as f:
                st.download_button("📥 Télécharger le rapport", f, file_name=f"rapport_analyse_{date_debut}_{date_fin}.docx")

            
            # Nettoyage des fichiers images générés
            #for img_path in [img_prod_path, img_etat_path, img_ev1_path, img_ev2_path]:
                #if os.path.exists(img_path):
                   # os.remove(img_path)




# Codé par Amboara RASOLOFOARIMANANA



                               


                                

                            
                            